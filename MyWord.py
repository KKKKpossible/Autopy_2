from docx import Document
import os
import tkinter
from tkinter import filedialog as fd  # import file dialog


class MyDocx:
    def __init__(self, path=None):
        if path is None:
            self.inst_doc = Document()
            self.inst_doc.add_heading('Document Title', 0)
            self.para = self.inst_doc.add_paragraph()
            self.dir_path = ''
        else:
            self.inst_doc = Document(path)
            self.dir_path = ''

    def save_docx(self):
        g_root = tkinter.Tk()
        g_root.withdraw()
        self.dir_path = fd.asksaveasfilename(parent=g_root,
                                             initialdir=os.getcwd(),
                                             initialfile='atr.docx',
                                             title='Save Word File',
                                             filetypes=[('word files', '*.docx'), ('all', '*.*')])
        g_root.quit()
        self.inst_doc.save(self.dir_path)

    def write_docx(self, text=''):
        pass


if __name__ == '__main__':
    f = open('test.docx', 'rb')
    document = Document(f)
    f.close()
